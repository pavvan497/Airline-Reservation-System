from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import BaseDocTemplate, Frame, Image as RLImage, PageBreak, PageTemplate, Paragraph, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
ASSET_DIR = OUT_DIR / "report_assets"
DOCX_PATH = OUT_DIR / "ARMS_IEEE_Project_Report.docx"
PDF_PATH = OUT_DIR / "ARMS_IEEE_Project_Report.pdf"


def get_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/timesbd.ttf" if bold else "C:/Windows/Fonts/times.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def create_architecture_diagram(path: Path) -> None:
    width, height = 1800, 1050
    img = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(img)
    title_font = get_font(42, bold=True)
    heading_font = get_font(30, bold=True)
    body_font = get_font(24)
    small_font = get_font(20)

    draw.text((90, 40), "Airline Reservation Management System Architecture", fill="#0d1b2a", font=title_font)

    boxes = [
        ((120, 180, 520, 510), "#dbeafe", "Presentation Layer", [
            "Thymeleaf templates",
            "Passenger pages",
            "Admin dashboard",
            "JavaScript interactions",
        ]),
        ((700, 180, 1100, 510), "#dcfce7", "Application Layer", [
            "PageController",
            "AuthenticationController",
            "BookingService",
            "PlaneService",
        ]),
        ((1280, 180, 1680, 510), "#fee2e2", "Data Layer", [
            "JPA repositories",
            "User / Booking / AirPlane",
            "MySQL persistence",
            "Seed data initializer",
        ]),
        ((430, 650, 850, 930), "#ede9fe", "Security Layer", [
            "Spring Security",
            "JWT token service",
            "Password encoding",
            "Role-based access",
        ]),
        ((960, 650, 1470, 930), "#fef3c7", "Business Rules", [
            "Route validation",
            "Seat availability check",
            "Price = km x seats x base fare",
            "Booking reference generation",
        ]),
    ]

    for (x1, y1, x2, y2), fill, heading, lines in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=28, fill=fill, outline="#334155", width=4)
        draw.text((x1 + 24, y1 + 20), heading, fill="#0f172a", font=heading_font)
        y = y1 + 85
        for line in lines:
            draw.text((x1 + 34, y), f"- {line}", fill="#1e293b", font=body_font)
            y += 48

    arrows = [
        ((520, 345), (700, 345)),
        ((1100, 345), (1280, 345)),
        ((850, 790), (960, 790)),
        ((900, 510), (900, 650)),
    ]
    for start, end in arrows:
        draw.line((start, end), fill="#2563eb", width=8)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 22, ey - 12), (ex - 22, ey + 12)], fill="#2563eb")

    draw.text((120, 980), "Fig. 1. Layered view of the ARMS web application and its key responsibilities.", fill="#334155", font=small_font)
    img.save(path)


def create_workflow_diagram(path: Path) -> None:
    width, height = 1800, 760
    img = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(img)
    title_font = get_font(40, bold=True)
    step_font = get_font(26, bold=True)
    body_font = get_font(21)
    caption_font = get_font(20)

    draw.text((80, 35), "Passenger Booking Workflow", fill="#0d1b2a", font=title_font)

    steps = [
        ("1", "Register / Login", "User account is created and authenticated with JWT-based security."),
        ("2", "Search Route", "Available flights are filtered by origin, destination, and travel date."),
        ("3", "Price Check", "System validates route and computes fare using distance and requested seats."),
        ("4", "Reserve Seats", "Booking record is saved and available seat count is reduced atomically."),
        ("5", "Confirmation", "Reference number, ticket number, date, and payment summary are shown."),
    ]

    x = 80
    top = 190
    box_w = 300
    gap = 42
    for idx, (num, heading, body) in enumerate(steps):
        x1 = x + idx * (box_w + gap)
        x2 = x1 + box_w
        fill = ["#dbeafe", "#dcfce7", "#fef3c7", "#fee2e2", "#ede9fe"][idx]
        draw.rounded_rectangle((x1, top, x2, top + 300), radius=25, fill=fill, outline="#334155", width=4)
        draw.ellipse((x1 + 108, top - 45, x1 + 192, top + 39), fill="#1d4ed8", outline="#1e3a8a", width=4)
        draw.text((x1 + 140, top - 25), num, fill="#ffffff", font=step_font, anchor="mm")
        draw.text((x1 + 20, top + 28), heading, fill="#0f172a", font=step_font)

        words = body.split()
        lines = []
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if draw.textlength(candidate, font=body_font) <= box_w - 40:
                current = candidate
            else:
                lines.append(current)
                current = word
        if current:
            lines.append(current)

        y = top + 90
        for line in lines:
            draw.text((x1 + 20, y), line, fill="#334155", font=body_font)
            y += 34

        if idx < len(steps) - 1:
            start = (x2 + 5, top + 150)
            end = (x2 + gap - 5, top + 150)
            draw.line((start, end), fill="#2563eb", width=8)
            draw.polygon([(end[0], end[1]), (end[0] - 24, end[1] - 14), (end[0] - 24, end[1] + 14)], fill="#2563eb")

    draw.text((80, 680), "Fig. 2. Sequential workflow followed by a passenger from authentication to final booking confirmation.", fill="#334155", font=caption_font)
    img.save(path)


def configure_page(section) -> None:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.gutter = Inches(0)
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)


def add_paragraph(document: Document, text: str, *, style=None, bold=False, italic=False, size=10, space_after=3, align=None):
    p = document.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    return p


def add_heading_like(document: Document, text: str, *, level: str) -> None:
    size = 11 if level == "main" else 10
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(5 if level == "main" else 2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)


def add_figure(document: Document, image_path: Path, caption: str, width: float) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    p.paragraph_format.space_after = Pt(2)
    add_paragraph(document, caption, italic=True, size=9, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_reference(document: Document, label: str, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.hanging_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"[{label}] {text}")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(9)


def build_document() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    ASSET_DIR.mkdir(exist_ok=True)
    arch_png = ASSET_DIR / "arms_architecture.png"
    workflow_png = ASSET_DIR / "arms_workflow.png"
    create_architecture_diagram(arch_png)
    create_workflow_diagram(workflow_png)

    doc = Document()
    section = doc.sections[0]
    configure_page(section)
    section.start_type = WD_SECTION.CONTINUOUS
    sect_pr = section._sectPr
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "720")
    sect_pr.append(cols)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(10)

    add_paragraph(
        doc,
        "Airline Reservation Management System:\nA Secure Spring Boot Web Application",
        bold=True,
        size=16,
        space_after=4,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        doc,
        "Project Team\nDepartment of Computer Engineering\nAcademic Project Report, 2025-2026",
        size=10,
        space_after=6,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Abstract- ")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run2 = p.add_run(
        "This paper presents the Airline Reservation Management System (ARMS), a web-based platform for flight discovery, passenger registration, fare estimation, reservation processing, and administrative monitoring. "
        "The application is implemented using Spring Boot 3, Thymeleaf, Spring Security, JPA, MySQL, and client-side JavaScript. "
        "ARMS maintains a clear separation between presentation, business, and persistence layers while supporting route-based flight search, distance-driven fare calculation, seat availability checks, booking confirmation generation, and dashboard-level operational summaries. "
        "The system uses JWT-backed authentication for account security and persists key entities such as users, flights, and bookings in a relational database. "
        "The resulting platform demonstrates how a modern Java web stack can be used to digitalize airline booking workflows in an accessible and modular manner."
    )
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(10)
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Index Terms- ")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run2 = p.add_run("Airline Reservation, Spring Boot, JWT Authentication, Thymeleaf, MySQL, Booking System, Admin Dashboard.")
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(10)
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_heading_like(doc, "I. INTRODUCTION", level="main")
    add_paragraph(doc, "Air transportation systems demand timely reservation handling, consistent seat inventory management, and secure access control for both passengers and administrators. Traditional manual or semi-digital reservation workflows often suffer from delayed updates, duplicate work, and limited visibility into real-time seat status. The Airline Reservation Management System addresses these limitations through a centralized web platform that allows passengers to register, browse routes, verify prices, and complete bookings while enabling administrators to monitor flights, passengers, revenue, and booking trends from a unified interface.")

    add_heading_like(doc, "A. Problem Statement", level="sub")
    add_paragraph(doc, "A flight booking platform must validate origin and destination pairs, protect user accounts, compute fares consistently, avoid overbooking, and provide immediate booking feedback. In many academic prototypes, these concerns are handled in isolation, resulting in fragmented flows and weak traceability across the reservation lifecycle. The core problem addressed by ARMS is the integration of secure authentication, dynamic route selection, accurate seat tracking, and operational visibility in one maintainable application.")

    add_heading_like(doc, "B. Motivation", level="sub")
    add_paragraph(doc, "The project is motivated by the need to model a realistic service workflow using enterprise Java technologies. Airline booking is a familiar but technically rich domain that combines data validation, transactional updates, user management, and dashboard-style reporting. Building ARMS provides hands-on experience with full-stack development while also demonstrating how core software engineering principles such as modularization, persistence, and role-aware access can be applied to a practical transport use case.")

    add_heading_like(doc, "C. Project Scope and Objectives", level="sub")
    add_paragraph(doc, "The scope of ARMS includes passenger registration, secure login, route browsing, seat-aware booking, fare estimation, booking confirmation, and administrative monitoring. The primary objectives are to: (i) provide a smooth booking experience for end users; (ii) maintain route and seat information in a relational database; (iii) secure identity flows through JWT-based authentication; (iv) support administrator functions for adding flights and viewing demand patterns; and (v) keep the codebase extensible for future features such as cancellation, payment gateways, and notification services.")

    add_heading_like(doc, "II. LITERATURE REVIEW", level="main")
    add_paragraph(doc, "Reservation systems are a standard case study in web engineering because they combine multi-user access, transactional integrity, and data-centric workflows. Prior studies on transportation booking platforms emphasize the importance of normalized relational data models, search usability, and security-first session handling. Within the Java ecosystem, Spring Boot is widely adopted for rapid backend development, while Thymeleaf remains a practical templating choice for server-rendered pages. JPA-based persistence and MySQL-backed storage are also common patterns for academic reservation systems due to their approachable development model and strong community support.")

    add_paragraph(doc, "The ARMS implementation aligns with these patterns by using entity-driven persistence for core business objects, controller-service layering for request handling, and a template-based interface for passengers and administrators. Unlike a purely static prototype, the project goes further by preserving live seat counts, assigning booking references, and exposing dashboard metrics such as total bookings, revenue, and route popularity.")

    add_heading_like(doc, "III. SYSTEM ARCHITECTURE AND METHODOLOGY", level="main")
    add_paragraph(doc, "ARMS follows a layered architecture that separates user interaction, business logic, and persistence concerns. This structure improves maintainability and makes it easier to evolve individual modules without destabilizing the overall reservation flow.")
    add_figure(doc, arch_png, "Fig. 1. Layered architecture of the Airline Reservation Management System.", 3.25)

    add_heading_like(doc, "A. Architectural Framework", level="sub")
    add_paragraph(doc, "The presentation layer consists of Thymeleaf templates under passenger and administrator views, styled with CSS and supported by JavaScript for client-side interactions. The application layer includes controllers for routing requests and services for implementing business rules such as price estimation, seat validation, and booking persistence. The data layer uses JPA repositories over MySQL for storing flights, users, and bookings. Security concerns are managed by Spring Security and JWT filters, allowing the system to attach identity information to protected operations.")

    add_heading_like(doc, "B. Execution Methodology", level="sub")
    add_paragraph(doc, "The workflow begins when a passenger registers or authenticates into the system. After login, the user selects an origin, destination, number of seats, and travel date. The booking service first rejects invalid requests, such as identical source and destination values or dates earlier than the current day. It then queries the flight repository for a matching route, calculates the fare using the configured base fare per kilometer, and checks whether sufficient seats remain. If validation succeeds, the booking record is stored, the flight inventory is reduced, and a confirmation response containing the booking reference and ticket number is generated.")
    add_figure(doc, workflow_png, "Fig. 2. Booking workflow executed by ARMS during reservation processing.", 3.25)

    add_heading_like(doc, "C. User Interface Design", level="sub")
    add_paragraph(doc, "The user-facing interface is organized into pages for home, route selection, payment, ticket display, booking confirmation, and booking history. On the administrative side, the dashboard summarizes revenue, booking counts, passenger totals, low-seat flights, and popular routes. This division enables each actor to access only the information relevant to their role while keeping the navigation model straightforward for an academic demonstration system.")

    add_heading_like(doc, "IV. IMPLEMENTATION DETAILS", level="main")
    add_heading_like(doc, "A. Technological Stack and Environment", level="sub")
    add_paragraph(doc, "The backend is implemented in Java 17 using Spring Boot 3.0.1 with dependencies for Spring Web, Spring Data JPA, Spring Security, Thymeleaf, Mail, JWT processing, and Lombok. MySQL serves as the persistent data store. Frontend pages are built with HTML templates, CSS, Bootstrap-based styling assets, and JavaScript. This combination provides a balanced stack for rapid development while preserving strong support for structured application design.")

    add_heading_like(doc, "B. Data Structures and Object Modeling", level="sub")
    add_paragraph(doc, "Three main entities define the data model. The User entity stores identity fields, encrypted passwords, and role information for both passengers and administrators. The AirPlane entity stores route endpoints, available seat count, and flight distance in kilometers. The Booking entity records the selected route, passenger email, seat count, travel date, total price, booking reference, and ticket number. These classes are mapped to relational tables and accessed through dedicated repository interfaces.")

    add_heading_like(doc, "C. Security and Authentication", level="sub")
    add_paragraph(doc, "Authentication is implemented through Spring Security with a JWT filter placed before the username-password authentication filter. During registration, passwords are encoded before persistence and each user is assigned either a passenger or administrator role. During login, credentials are authenticated through the configured authentication manager, and the JWT service generates a token for subsequent requests. This design supports stateless security handling and keeps protected business operations tied to the active user identity.")

    add_heading_like(doc, "D. Core Booking Logic", level="sub")
    add_paragraph(doc, "The BookingService acts as the primary business module for reservation processing. It calculates price using a base fare of 12.0 per kilometer multiplied by the requested number of seats, validates that the travel date is not in the past, checks that source and destination are different, and verifies that a matching flight exists. When a booking is confirmed, the service creates a unique booking reference and ticket number, stores the transaction, and decrements the corresponding flight's available seat count. This logic prevents overbooking while maintaining a traceable reservation record.")

    add_heading_like(doc, "E. Administrative Features", level="sub")
    add_paragraph(doc, "The administrative module exposes dashboard-level summaries derived from repository data. It reports total bookings, total revenue, total passengers, total flights, flights with low remaining seat counts, and a ranked list of heavily booked routes. Administrators may also add new flights through a dedicated form that captures route endpoints, seat capacity, and route distance. A demo data initializer seeds representative flight pairs and baseline user accounts to simplify testing and presentation.")

    add_heading_like(doc, "V. RESULTS AND DISCUSSION", level="main")
    add_paragraph(doc, "Functional review of the application shows that the system successfully supports the main reservation lifecycle from user onboarding to booking confirmation. Route lists are populated from stored flight records, seat inventory updates after booking, and generated identifiers improve traceability of each transaction. The administrator views further convert raw repository data into immediately useful indicators for operational monitoring.")

    add_heading_like(doc, "A. Functional Outcomes", level="sub")
    add_paragraph(doc, "The completed application demonstrates that a server-rendered Spring Boot solution can deliver a coherent airline reservation experience without requiring a separate single-page frontend. Users are able to discover routes, estimate price, reserve seats, and revisit booking data, while administrators can inspect route pressure and revenue summaries. The modular code organization also supports incremental addition of new service logic.")

    add_heading_like(doc, "B. Module Summary", level="sub")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Module", "Primary Responsibility", "Key Technologies", "Observed Benefit"]
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        cell.text = text
        set_cell_shading(cell, "D9EAF7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(9)
    rows = [
        ("Authentication", "Register users, encode passwords, generate JWT tokens", "Spring Security, JWT", "Secures account access and role handling"),
        ("Flight Management", "Store route, distance, and available seat data", "JPA, MySQL", "Keeps route inventory centralized"),
        ("Booking Engine", "Validate travel data, compute fare, save reservation", "Spring Services", "Supports accurate and traceable booking"),
        ("Admin Dashboard", "Summarize bookings, revenue, and low-seat flights", "Thymeleaf, Repository queries", "Improves operational visibility"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9)

    add_heading_like(doc, "VI. CONCLUSION AND FUTURE WORK", level="main")
    add_paragraph(doc, "The Airline Reservation Management System demonstrates a practical application of secure full-stack Java development to a real-world reservation domain. By combining Spring Boot, MySQL, Thymeleaf, and JWT-backed security, the project achieves a functional workflow for registration, route discovery, fare estimation, seat-aware booking, and administrative monitoring. The resulting system is suitable as both an academic project and a foundation for further product-level refinement.")
    add_paragraph(doc, "Future improvements may include online payment gateway integration, email notifications, ticket cancellation with seat restoration, route schedules with departure times, stronger transactional safeguards around concurrent booking, and analytics dashboards with chart-based visualizations.")

    add_heading_like(doc, "REFERENCES", level="main")
    add_reference(doc, "1", "C. Richardson, Microservices Patterns. Manning Publications, 2018.")
    add_reference(doc, "2", "C. Walls, Spring Boot in Action. Manning Publications, 2016.")
    add_reference(doc, "3", "R. Johnson et al., Professional Java Development with the Spring Framework. Wrox, 2005.")
    add_reference(doc, "4", "Oracle, 'The Java Language Specification, Java SE 17 Edition,' 2021.")
    add_reference(doc, "5", "P. DuBois, MySQL. Addison-Wesley, 2021.")

    doc.save(DOCX_PATH)


def build_pdf() -> None:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleIEEE",
        parent=styles["Title"],
        fontName="Times-Bold",
        fontSize=18,
        leading=20,
        alignment=TA_CENTER,
        spaceAfter=8,
    )
    author_style = ParagraphStyle(
        "AuthorIEEE",
        parent=styles["Normal"],
        fontName="Times-Roman",
        fontSize=10,
        leading=12,
        alignment=TA_CENTER,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "BodyIEEE",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=9,
        leading=11,
        alignment=TA_JUSTIFY,
        spaceAfter=5,
    )
    heading_style = ParagraphStyle(
        "HeadingIEEE",
        parent=body_style,
        fontName="Times-Bold",
        fontSize=10,
        leading=12,
        spaceBefore=6,
        spaceAfter=3,
    )
    subheading_style = ParagraphStyle(
        "SubHeadingIEEE",
        parent=body_style,
        fontName="Times-Bold",
        fontSize=9,
        leading=11,
        spaceBefore=4,
        spaceAfter=2,
    )
    abstract_label = ParagraphStyle(
        "AbstractLabel",
        parent=body_style,
        fontName="Times-Bold",
    )
    caption_style = ParagraphStyle(
        "CaptionIEEE",
        parent=body_style,
        fontSize=8,
        leading=9,
        alignment=TA_CENTER,
        spaceAfter=5,
    )
    ref_style = ParagraphStyle(
        "RefIEEE",
        parent=body_style,
        fontSize=8.5,
        leading=10,
        leftIndent=12,
        firstLineIndent=-12,
        spaceAfter=2,
    )

    doc = BaseDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
    )
    col_gap = 0.22 * inch
    frame_width = (doc.width - col_gap) / 2
    frame_height = doc.height
    frames = [
        Frame(doc.leftMargin, doc.bottomMargin, frame_width, frame_height, id="left"),
        Frame(doc.leftMargin + frame_width + col_gap, doc.bottomMargin, frame_width, frame_height, id="right"),
    ]
    doc.addPageTemplates([PageTemplate(id="TwoCol", frames=frames)])

    story = []
    story.append(Paragraph("Airline Reservation Management System:<br/>A Secure Spring Boot Web Application", title_style))
    story.append(Paragraph("Project Team<br/>Department of Computer Engineering<br/>Academic Project Report, 2025-2026", author_style))
    story.append(Paragraph("<b>Abstract-</b> This paper presents the Airline Reservation Management System (ARMS), a web-based platform for flight discovery, passenger registration, fare estimation, reservation processing, and administrative monitoring. The application is implemented using Spring Boot 3, Thymeleaf, Spring Security, JPA, MySQL, and client-side JavaScript. ARMS maintains a clear separation between presentation, business, and persistence layers while supporting route-based flight search, distance-driven fare calculation, seat availability checks, booking confirmation generation, and dashboard-level operational summaries. The system uses JWT-backed authentication for account security and persists key entities such as users, flights, and bookings in a relational database. The resulting platform demonstrates how a modern Java web stack can be used to digitalize airline booking workflows in an accessible and modular manner.", body_style))
    story.append(Paragraph("<b>Index Terms-</b> Airline Reservation, Spring Boot, JWT Authentication, Thymeleaf, MySQL, Booking System, Admin Dashboard.", body_style))

    sections = [
        ("I. INTRODUCTION", [
            ("body", "Air transportation systems demand timely reservation handling, consistent seat inventory management, and secure access control for both passengers and administrators. Traditional manual or semi-digital reservation workflows often suffer from delayed updates, duplicate work, and limited visibility into real-time seat status. The Airline Reservation Management System addresses these limitations through a centralized web platform that allows passengers to register, browse routes, verify prices, and complete bookings while enabling administrators to monitor flights, passengers, revenue, and booking trends from a unified interface."),
            ("sub", "A. Problem Statement"),
            ("body", "A flight booking platform must validate origin and destination pairs, protect user accounts, compute fares consistently, avoid overbooking, and provide immediate booking feedback. In many academic prototypes, these concerns are handled in isolation, resulting in fragmented flows and weak traceability across the reservation lifecycle. The core problem addressed by ARMS is the integration of secure authentication, dynamic route selection, accurate seat tracking, and operational visibility in one maintainable application."),
            ("sub", "B. Motivation"),
            ("body", "The project is motivated by the need to model a realistic service workflow using enterprise Java technologies. Airline booking is a familiar but technically rich domain that combines data validation, transactional updates, user management, and dashboard-style reporting. Building ARMS provides hands-on experience with full-stack development while also demonstrating how core software engineering principles such as modularization, persistence, and role-aware access can be applied to a practical transport use case."),
            ("sub", "C. Project Scope and Objectives"),
            ("body", "The scope of ARMS includes passenger registration, secure login, route browsing, seat-aware booking, fare estimation, booking confirmation, and administrative monitoring. The primary objectives are to: (i) provide a smooth booking experience for end users; (ii) maintain route and seat information in a relational database; (iii) secure identity flows through JWT-based authentication; (iv) support administrator functions for adding flights and viewing demand patterns; and (v) keep the codebase extensible for future features such as cancellation, payment gateways, and notification services."),
        ]),
        ("II. LITERATURE REVIEW", [
            ("body", "Reservation systems are a standard case study in web engineering because they combine multi-user access, transactional integrity, and data-centric workflows. Prior studies on transportation booking platforms emphasize the importance of normalized relational data models, search usability, and security-first session handling. Within the Java ecosystem, Spring Boot is widely adopted for rapid backend development, while Thymeleaf remains a practical templating choice for server-rendered pages. JPA-based persistence and MySQL-backed storage are also common patterns for academic reservation systems due to their approachable development model and strong community support."),
            ("body", "The ARMS implementation aligns with these patterns by using entity-driven persistence for core business objects, controller-service layering for request handling, and a template-based interface for passengers and administrators. Unlike a purely static prototype, the project goes further by preserving live seat counts, assigning booking references, and exposing dashboard metrics such as total bookings, revenue, and route popularity."),
        ]),
        ("III. SYSTEM ARCHITECTURE AND METHODOLOGY", [
            ("body", "ARMS follows a layered architecture that separates user interaction, business logic, and persistence concerns. This structure improves maintainability and makes it easier to evolve individual modules without destabilizing the overall reservation flow."),
            ("sub", "A. Architectural Framework"),
            ("body", "The presentation layer consists of Thymeleaf templates under passenger and administrator views, styled with CSS and supported by JavaScript for client-side interactions. The application layer includes controllers for routing requests and services for implementing business rules such as price estimation, seat validation, and booking persistence. The data layer uses JPA repositories over MySQL for storing flights, users, and bookings. Security concerns are managed by Spring Security and JWT filters, allowing the system to attach identity information to protected operations."),
            ("sub", "B. Execution Methodology"),
            ("body", "The workflow begins when a passenger registers or authenticates into the system. After login, the user selects an origin, destination, number of seats, and travel date. The booking service first rejects invalid requests, such as identical source and destination values or dates earlier than the current day. It then queries the flight repository for a matching route, calculates the fare using the configured base fare per kilometer, and checks whether sufficient seats remain. If validation succeeds, the booking record is stored, the flight inventory is reduced, and a confirmation response containing the booking reference and ticket number is generated."),
            ("sub", "C. User Interface Design"),
            ("body", "The user-facing interface is organized into pages for home, route selection, payment, ticket display, booking confirmation, and booking history. On the administrative side, the dashboard summarizes revenue, booking counts, passenger totals, low-seat flights, and popular routes. This division enables each actor to access only the information relevant to their role while keeping the navigation model straightforward for an academic demonstration system."),
        ]),
        ("IV. IMPLEMENTATION DETAILS", [
            ("sub", "A. Technological Stack and Environment"),
            ("body", "The backend is implemented in Java 17 using Spring Boot 3.0.1 with dependencies for Spring Web, Spring Data JPA, Spring Security, Thymeleaf, Mail, JWT processing, and Lombok. MySQL serves as the persistent data store. Frontend pages are built with HTML templates, CSS, Bootstrap-based styling assets, and JavaScript. This combination provides a balanced stack for rapid development while preserving strong support for structured application design."),
            ("sub", "B. Data Structures and Object Modeling"),
            ("body", "Three main entities define the data model. The User entity stores identity fields, encrypted passwords, and role information for both passengers and administrators. The AirPlane entity stores route endpoints, available seat count, and flight distance in kilometers. The Booking entity records the selected route, passenger email, seat count, travel date, total price, booking reference, and ticket number. These classes are mapped to relational tables and accessed through dedicated repository interfaces."),
            ("sub", "C. Security and Authentication"),
            ("body", "Authentication is implemented through Spring Security with a JWT filter placed before the username-password authentication filter. During registration, passwords are encoded before persistence and each user is assigned either a passenger or administrator role. During login, credentials are authenticated through the configured authentication manager, and the JWT service generates a token for subsequent requests. This design supports stateless security handling and keeps protected business operations tied to the active user identity."),
            ("sub", "D. Core Booking Logic"),
            ("body", "The BookingService acts as the primary business module for reservation processing. It calculates price using a base fare of 12.0 per kilometer multiplied by the requested number of seats, validates that the travel date is not in the past, checks that source and destination are different, and verifies that a matching flight exists. When a booking is confirmed, the service creates a unique booking reference and ticket number, stores the transaction, and decrements the corresponding flight's available seat count. This logic prevents overbooking while maintaining a traceable reservation record."),
            ("sub", "E. Administrative Features"),
            ("body", "The administrative module exposes dashboard-level summaries derived from repository data. It reports total bookings, total revenue, total passengers, total flights, flights with low remaining seat counts, and a ranked list of heavily booked routes. Administrators may also add new flights through a dedicated form that captures route endpoints, seat capacity, and route distance. A demo data initializer seeds representative flight pairs and baseline user accounts to simplify testing and presentation."),
        ]),
        ("V. RESULTS AND DISCUSSION", [
            ("body", "Functional review of the application shows that the system successfully supports the main reservation lifecycle from user onboarding to booking confirmation. Route lists are populated from stored flight records, seat inventory updates after booking, and generated identifiers improve traceability of each transaction. The administrator views further convert raw repository data into immediately useful indicators for operational monitoring."),
            ("sub", "A. Functional Outcomes"),
            ("body", "The completed application demonstrates that a server-rendered Spring Boot solution can deliver a coherent airline reservation experience without requiring a separate single-page frontend. Users are able to discover routes, estimate price, reserve seats, and revisit booking data, while administrators can inspect route pressure and revenue summaries. The modular code organization also supports incremental addition of new service logic."),
        ]),
        ("VI. CONCLUSION AND FUTURE WORK", [
            ("body", "The Airline Reservation Management System demonstrates a practical application of secure full-stack Java development to a real-world reservation domain. By combining Spring Boot, MySQL, Thymeleaf, and JWT-backed security, the project achieves a functional workflow for registration, route discovery, fare estimation, seat-aware booking, and administrative monitoring. The resulting system is suitable as both an academic project and a foundation for further product-level refinement."),
            ("body", "Future improvements may include online payment gateway integration, email notifications, ticket cancellation with seat restoration, route schedules with departure times, stronger transactional safeguards around concurrent booking, and analytics dashboards with chart-based visualizations."),
        ]),
    ]

    for heading, parts in sections:
        story.append(Paragraph(heading, heading_style))
        for kind, text in parts:
            story.append(Paragraph(text, subheading_style if kind == "sub" else body_style))
        if heading == "III. SYSTEM ARCHITECTURE AND METHODOLOGY":
            story.append(RLImage(str(ASSET_DIR / "arms_architecture.png"), width=3.0 * inch, height=1.75 * inch))
            story.append(Paragraph("Fig. 1. Layered architecture of the Airline Reservation Management System.", caption_style))
            story.append(RLImage(str(ASSET_DIR / "arms_workflow.png"), width=3.1 * inch, height=1.32 * inch))
            story.append(Paragraph("Fig. 2. Booking workflow executed by ARMS during reservation processing.", caption_style))
        if heading == "V. RESULTS AND DISCUSSION":
            table_data = [
                ["Module", "Primary Responsibility", "Key Technologies", "Observed Benefit"],
                ["Authentication", "Register users and generate tokens", "Spring Security, JWT", "Secures access"],
                ["Flight Management", "Store route and seat data", "JPA, MySQL", "Centralized inventory"],
                ["Booking Engine", "Validate and save reservations", "Spring Services", "Seat-aware booking"],
                ["Admin Dashboard", "Summarize routes and revenue", "Thymeleaf, repositories", "Operational visibility"],
            ]
            table = Table(table_data, colWidths=[0.95 * inch, 1.35 * inch, 1.1 * inch, 1.05 * inch])
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9eaf7")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Times-Roman"),
                ("FONTSIZE", (0, 0), (-1, -1), 7.4),
                ("LEADING", (0, 0), (-1, -1), 8.4),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(table)
            story.append(Spacer(1, 6))

    story.append(Paragraph("REFERENCES", heading_style))
    refs = [
        "[1] C. Richardson, Microservices Patterns. Manning Publications, 2018.",
        "[2] C. Walls, Spring Boot in Action. Manning Publications, 2016.",
        "[3] R. Johnson et al., Professional Java Development with the Spring Framework. Wrox, 2005.",
        "[4] Oracle, 'The Java Language Specification, Java SE 17 Edition,' 2021.",
        "[5] P. DuBois, MySQL. Addison-Wesley, 2021.",
    ]
    for ref in refs:
        story.append(Paragraph(ref, ref_style))

    doc.build(story)


if __name__ == "__main__":
    build_document()
    build_pdf()
